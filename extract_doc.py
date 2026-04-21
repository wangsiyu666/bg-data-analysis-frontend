import io, sys
from docx import Document

doc_path = r"C:\work\baoge\backend\telecompass智能体前后端接口说明.docx"
doc = Document(doc_path)

out = io.open(r"C:\work\baoge\doc_extracted.txt", "w", encoding="utf-8")
out.write("=" * 80 + "\n")
out.write("PARAGRAPHS\n")
out.write("=" * 80 + "\n")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if t.strip():
        out.write(f"[P{i}] {t}\n")

out.write("\n" + "=" * 80 + "\n")
out.write("TABLES\n")
out.write("=" * 80 + "\n")
for ti, table in enumerate(doc.tables):
    out.write(f"\n--- Table {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---\n")
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
        out.write(f"  R{ri}: " + " || ".join(cells) + "\n")
out.close()
print("done")
